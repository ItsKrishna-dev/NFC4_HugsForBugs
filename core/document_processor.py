"""
core/document_processor.py
A lightweight document-processing pipeline that
1. extracts text from PDF / DOCX / TXT,
2. splits it into overlapping chunks,
3. creates TF-IDF embeddings,
4. stores everything in SQLite.

Compatible with Python 3.10+ and the simplified
requirements (PyPDF2, python-docx, scikit-learn, numpy).
"""

from __future__ import annotations

import hashlib
import logging
import re
from datetime import datetime
from pathlib import Path
from typing import Dict, List

import numpy as np
import PyPDF2
from docx import Document as DocxDocument
from sklearn.feature_extraction.text import TfidfVectorizer

from config.settings import CHUNK_OVERLAP, CHUNK_SIZE, SUPPORTED_FORMATS
from utils.database import get_connection
from utils.text_utils import clean_text

# --------------------------------------------------------------------------- #
#  logging setup
# --------------------------------------------------------------------------- #
logging.basicConfig(level=logging.INFO, format="%(levelname)s | %(message)s")
log = logging.getLogger(__name__)


# --------------------------------------------------------------------------- #
#  helpers: splitter & embedder
# --------------------------------------------------------------------------- #
class SimpleTextSplitter:
    """Chunk text with fixed size and optional overlap."""

    def __init__(self, chunk_size: int = 1_000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def split_text(self, text: str) -> List[str]:
        if len(text) <= self.chunk_size:
            return [text]

        chunks, start = [], 0
        while start < len(text):
            end = start + self.chunk_size
            chunk = text[start:end]

            # try to break at a sentence boundary for nicer chunks
            if end < len(text):
                last_break = max(chunk.rfind("."), chunk.rfind("\n"))
                if last_break > start + self.chunk_size // 2:
                    end = start + last_break + 1
                    chunk = text[start:end]

            chunks.append(chunk.strip())
            start = end - self.chunk_overlap

        return [c for c in chunks if c.strip()]


class SimpleEmbedder:
    """TF-IDF embeddings (keeps everything local & free)."""

    def __init__(self):
        self.vectorizer = TfidfVectorizer(max_features=384, stop_words="english")
        self._fitted = False

    def encode(self, texts: List[str]) -> np.ndarray:
        if not self._fitted:
            mat = self.vectorizer.fit_transform(texts)
            self._fitted = True
        else:
            mat = self.vectorizer.transform(texts)
        return mat.toarray()


# --------------------------------------------------------------------------- #
#  DocumentProcessor
# --------------------------------------------------------------------------- #
class DocumentProcessor:
    def __init__(self) -> None:
        self.splitter = SimpleTextSplitter(CHUNK_SIZE, CHUNK_OVERLAP)
        self.embedder = SimpleEmbedder()
        self.conn = get_connection()
        self._init_db()

    # --------------------------- DB schema --------------------------------- #
    def _init_db(self) -> None:
        """Create tables with the **file_hash** column (no ‘hash’ anywhere)."""
        self.conn.execute(
            """
            CREATE TABLE IF NOT EXISTS documents (
                id            INTEGER PRIMARY KEY AUTOINCREMENT,
                file_hash     TEXT UNIQUE,
                filename      TEXT,
                processed_at  TIMESTAMP,          -- <── use processed_at
                total_chunks  INTEGER,
                total_chars   INTEGER
            )
            """
        )
        self.conn.execute(
            """
            CREATE TABLE IF NOT EXISTS chunks (
                id           INTEGER PRIMARY KEY AUTOINCREMENT,
                doc_id       INTEGER,
                idx          INTEGER,
                text         TEXT,
                embedding    BLOB,
                FOREIGN KEY(doc_id) REFERENCES documents(id)
            )
            """
        )
        self.conn.commit()

    # ---------------------- internal utilities ----------------------------- #
    @staticmethod
    def _sha256(b: bytes) -> str:
        return hashlib.sha256(b).hexdigest()

    def _is_cached(self, file_hash: str) -> bool:
        row = self.conn.execute(
            "SELECT 1 FROM documents WHERE file_hash = ?", (file_hash,)
        ).fetchone()
        return row is not None

    # ---------------------- text extraction -------------------------------- #
    @staticmethod
    def _pdf_text(path: Path) -> str:
        text = ""
        with path.open("rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                try:
                    text += page.extract_text() + "\n"
                except Exception as exc:  # noqa: BLE001
                    log.warning("PDF page skipped: %s", exc)
        return text

    @staticmethod
    def _docx_text(path: Path) -> str:
        doc = DocxDocument(path)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        # include table contents
        for table in doc.tables:
            for row in table.rows:
                text += "\n" + " | ".join(cell.text for cell in row.cells)
        return text

    @staticmethod
    def _txt_text(path: Path) -> str:
        encodings = ["utf-8", "latin-1", "cp1252", "utf-16"]
        for enc in encodings:
            try:
                return path.read_text(encoding=enc)
            except UnicodeDecodeError:
                continue
        raise UnicodeDecodeError("Could not decode text file")

    def _extract_text(self, path: Path) -> str:
        ext = path.suffix.lower()
        if ext == ".pdf":
            return self._pdf_text(path)
        if ext == ".docx":
            return self._docx_text(path)
        if ext in {".txt", ".md", ".rtf"}:
            return self._txt_text(path)
        raise ValueError(f"Unsupported file type: {ext}")

    def extract_sections(self, path: Path) -> list:
        """
        Extract sections from a document using style-aware methods if available.
        Returns a list of (section_title, section_content) tuples.
        """
        ext = path.suffix.lower()
        if ext == ".docx":
            return self._docx_sections(path)
        # PDF and TXT: fallback to text-based detection for now
        text = self._extract_text(path)
        from core.summarizer import Summarizer
        return Summarizer()._detect_sections_advanced(text)

    @staticmethod
    def _docx_sections(path: Path) -> list:
        """
        Extract sections from a DOCX file using heading styles and bold/large text cues.
        Returns a list of (section_title, section_content) tuples.
        """
        from docx.shared import Pt
        doc = DocxDocument(path)
        sections = []
        current_title = "Introduction"
        current_content = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = para.style.name.lower()
            is_heading = 'heading' in style
            is_bold = any(run.bold for run in para.runs if run.text.strip())
            # Try to get font size (may not always be available)
            font_sizes = [run.font.size.pt for run in para.runs if run.font.size is not None]
            avg_font_size = sum(font_sizes) / len(font_sizes) if font_sizes else 0
            # Heuristic: treat as section title if heading, or bold+large font, or short and capitalized
            if is_heading or (is_bold and avg_font_size >= 13) or (len(text) < 80 and text.istitle()):
                if current_content:
                    sections.append((current_title, "\n".join(current_content)))
                current_title = text
                current_content = []
            else:
                current_content.append(text)
        if current_content:
            sections.append((current_title, "\n".join(current_content)))
        return sections

    # ---------------------------- main API --------------------------------- #
    def process(self, path: Path, force_reprocess: bool = False) -> Dict:
        """Full pipeline: extract → clean → chunk → embed → persist."""
        if path.suffix.lower() not in SUPPORTED_FORMATS:
            raise ValueError("Unsupported file type")

        file_bytes = path.read_bytes()
        file_hash = self._sha256(file_bytes)

        if not force_reprocess and self._is_cached(file_hash):
            log.info("Cached - skipping %s", path.name)
            return {"status": "cached", "file_hash": file_hash}

        raw_text = self._extract_text(path)
        if not raw_text:
            raise ValueError("No text extracted from document")

        cleaned = clean_text(raw_text)
        chunks = self.splitter.split_text(cleaned)
        if not chunks:
            raise ValueError("Splitter returned 0 chunks")

        embeddings = self.embedder.encode(chunks)

        # ----- persist to DB ----- #
        cur = self.conn.execute(
            """
            INSERT INTO documents (file_hash, filename, processed_at,
                                   total_chunks, total_chars)
            VALUES (?, ?, ?, ?, ?)
            """,
            (file_hash,
             path.name,
             datetime.now().isoformat(),  # <── processed_at
             len(chunks),
             len(cleaned)),
        )
        doc_id = cur.lastrowid

        for idx, (chunk, emb) in enumerate(zip(chunks, embeddings)):
            self.conn.execute(
                """
                INSERT INTO chunks (doc_id, idx, text, embedding)
                VALUES (?, ?, ?, ?)
                """,
                (doc_id, idx, chunk, emb.tobytes()),
            )
        self.conn.commit()

        log.info("Processed %s | %s chunks", path.name, len(chunks))
        return {
            "status": "processed",
            "file_hash": file_hash,
            "doc_id": doc_id,
            "chunks": len(chunks),
        }

    # -------------------------- stats helper -------------------------------- #
    def get_processing_stats(self) -> Dict:
        c = self.conn.execute(
            """
            SELECT
              COUNT(*)               AS n_docs,
              COALESCE(SUM(total_chunks), 0),
              COALESCE(SUM(total_chars), 0),
              COALESCE(AVG(total_chunks), 0)
            FROM documents
            """
        ).fetchone()

        return {
            "total_documents": c[0],
            "total_chunks": c[1],
            "total_characters": c[2],
            "avg_chunks_per_doc": round(c[3], 2),
        }
